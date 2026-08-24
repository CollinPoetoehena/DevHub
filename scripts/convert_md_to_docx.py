#!/usr/bin/env python3
"""
convert_md_to_docx.py — Convert a Markdown file to a Word (.docx) document.

Parses common Markdown constructs (headings, tables, code blocks, inline formatting,
lists, blockquotes, horizontal rules) and produces a styled Word document using
python-docx. The output file is placed next to the input file with a .docx extension.

Requirements:
    pip install python-docx

Usage:
    python3 convert_md_to_docx.py <input.md> [output.docx]

    If output.docx is omitted, the output file is the input filename with .docx extension.

Examples:
    python3 convert_md_to_docx.py ../homelab/docs/reference/network/Subnets_and_IP_Addresses.md /mnt/c/Users/poeto501/Downloads/Subnets_and_IP_Addresses.docx
    python3 convert_md_to_docx.py ../README.md
"""

import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print(
        "Error: python-docx is not installed.\n"
        "Install it with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# ── Styling helpers ──────────────────────────────────────────────────────────

def _set_cell_shading(cell, color: str):
    """Set background colour for a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    tc_pr.append(shading)


def _add_run_shading(run, color: str):
    """Apply character-level background shading to a run (for inline code)."""
    r_pr = run._element.get_or_add_rPr()
    shading = r_pr.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    r_pr.append(shading)


# ── Inline formatting ────────────────────────────────────────────────────────

_INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")


def _process_inline(paragraph, text: str):
    """Parse inline Markdown (code, bold, italic, links) into Word runs."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue

        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

        elif part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            for sub in re.split(r"(`[^`]+`)", inner):
                if sub.startswith("`") and sub.endswith("`"):
                    run = paragraph.add_run(sub[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
                    run.bold = True
                else:
                    run = paragraph.add_run(sub)
                    run.bold = True

        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True

        elif part.startswith("[") and "](" in part:
            link_text = part[1 : part.index("]")]
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
            run.underline = True

        else:
            paragraph.add_run(part)


# ── Block-level elements ─────────────────────────────────────────────────────

def _add_code_block(doc, text: str):
    """Render a fenced code block as monospaced paragraphs with a grey background."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.style = doc.styles["No Spacing"]
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        _add_run_shading(run, "F0F0F0")


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    """Extract rows from pipe-delimited Markdown table lines, skipping separator rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line[1:-1].split("|")]
            if all(re.match(r"^[-:\s]+$", c) for c in cells):
                continue
            rows.append(cells)
    return rows


def _add_table(doc, rows: list[list[str]]):
    """Insert a styled Word table from parsed Markdown rows."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= len(table.rows[i].cells):
                continue
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _process_inline(p, cell_text)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.font.size = Pt(8.5)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                _set_cell_shading(cell, "D9E2F3")


def _flush_blockquote(doc, lines: list[str]):
    """Write accumulated blockquote lines as a single grey-italic paragraph."""
    if not lines:
        return
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    _process_inline(p, " ".join(lines))
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True


def _add_horizontal_rule(doc):
    """Insert a thin horizontal line as a paragraph bottom-border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn("w:pBdr"), {})
    bottom = p_bdr.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "AAAAAA"},
    )
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ── Main conversion ──────────────────────────────────────────────────────────

def convert(input_path: str, output_path: str):
    """Read a Markdown file and write a formatted Word document."""
    with open(input_path, "r", encoding="utf-8") as fh:
        content = fh.read()

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Heading colours
    for level in range(1, 5):
        doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    lines = content.split("\n")
    idx = 0
    in_code = False
    code_lines: list[str] = []
    bq_lines: list[str] = []

    while idx < len(lines):
        line = lines[idx]

        # ── Fenced code blocks ───────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, "\n".join(code_lines))
                code_lines.clear()
                in_code = False
            else:
                _flush_blockquote(doc, bq_lines)
                bq_lines.clear()
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        # ── Tables ───────────────────────────────────────────────────────
        if (
            line.strip().startswith("|")
            and idx + 1 < len(lines)
            and lines[idx + 1].strip().startswith("|")
        ):
            _flush_blockquote(doc, bq_lines)
            bq_lines.clear()
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            _add_table(doc, _parse_table_rows(table_lines))
            doc.add_paragraph()
            continue

        # ── Blockquotes ──────────────────────────────────────────────────
        if line.strip().startswith("> "):
            text = line.strip()[2:].strip().replace("**Note:**", "Note:", 1)
            bq_lines.append(text)
            idx += 1
            continue
        if bq_lines and line.strip() == "":
            _flush_blockquote(doc, bq_lines)
            bq_lines.clear()
            idx += 1
            continue

        # ── Horizontal rules ────────────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            _flush_blockquote(doc, bq_lines)
            bq_lines.clear()
            _add_horizontal_rule(doc)
            idx += 1
            continue

        # ── Headings ─────────────────────────────────────────────────────
        heading_m = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_m:
            level = len(heading_m.group(1))
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_m.group(2).strip())
            doc.add_heading(text, level=min(level, 4))
            idx += 1
            continue

        # ── Bullet list items ────────────────────────────────────────────
        bullet_m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet_m:
            indent_level = len(bullet_m.group(1)) // 4
            text = bullet_m.group(2).strip()
            j = idx + 1
            while j < len(lines):
                nxt = lines[j]
                if (
                    nxt.strip()
                    and not re.match(r"^\s*[-*]\s+", nxt)
                    and not nxt.strip().startswith(("#", "```", "|", ">"))
                    and len(nxt) - len(nxt.lstrip()) > len(bullet_m.group(1))
                ):
                    text += " " + nxt.strip()
                    j += 1
                else:
                    break
            p = doc.add_paragraph(style="List Bullet")
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
            _process_inline(p, text)
            idx = j
            continue

        # ── Numbered list items ──────────────────────────────────────────
        num_m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if num_m:
            p = doc.add_paragraph(style="List Number")
            _process_inline(p, num_m.group(2).strip())
            idx += 1
            continue

        # ── Blank lines ─────────────────────────────────────────────────
        if line.strip() == "":
            idx += 1
            continue

        # ── Regular paragraphs ───────────────────────────────────────────
        text = line.strip()
        j = idx + 1
        while j < len(lines):
            nxt = lines[j]
            if (
                nxt.strip()
                and not nxt.strip().startswith(("#", "```", "|", ">", "- ", "* "))
                and not re.match(r"^\d+\.\s+", nxt.strip())
                and nxt.strip() not in ("---", "***", "___")
            ):
                text += " " + nxt.strip()
                j += 1
            else:
                break
        p = doc.add_paragraph()
        _process_inline(p, text)
        idx = j

    # Flush any trailing blockquote
    _flush_blockquote(doc, bq_lines)

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── CLI entry point ──────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert a Markdown file to a Word (.docx) document.",
    )
    parser.add_argument("input", help="Path to the input Markdown file.")
    parser.add_argument(
        "output",
        nargs="?",
        default=None,
        help="Path for the output .docx file (default: same name as input with .docx extension).",
    )
    args = parser.parse_args()

    input_path = os.path.abspath(args.input)
    if not os.path.isfile(input_path):
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if args.output:
        output_path = os.path.abspath(args.output)
    else:
        output_path = os.path.splitext(input_path)[0] + ".docx"

    convert(input_path, output_path)


if __name__ == "__main__":
    main()
